import argparse
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from hotel_app.ml.models import MODEL_REGISTRY

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Arial'
        if level == 1:
            run.font.color.rgb = RGBColor(15, 61, 86) # Dark teal
        elif level == 2:
            run.font.color.rgb = RGBColor(30, 100, 130)

def add_paragraph(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def add_image(doc, image_path, width_inches=6.0, caption=""):
    if Path(image_path).exists():
        doc.add_picture(image_path, width=Inches(width_inches))
        if caption:
            p = add_paragraph(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            doc.add_paragraph("\n")

def add_code_block(doc, code):
    # Create a light gray background paragraph for code if possible, or just a monospaced block
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)
    p.paragraph_format.left_indent = Inches(0.2)
    doc.add_paragraph("\n")

def build_report(output_file="Hotel_Cancellation_Final_Report.docx", team_members=None):
    if team_members is None:
        team_members = ["[Insert Team Member 1 Name]", "[Insert Team Member 2 Name]", "[Insert Team Member 3 Name]"]

    doc = Document()

    # --- Title Page ---
    title = doc.add_heading("Hotel Booking Cancellation Prediction", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Machine Learning Final Project Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].bold = True
    doc.add_paragraph("\n")

    add_image(doc, "artifacts/plots/guest_segmentation.png", width_inches=5.5, caption="Guest Segmentation Visualization (PCA & K-Means)")
    
    doc.add_paragraph("\n\n")
    team_heading = doc.add_heading("Team Members", level=2)
    team_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for member in team_members:
        p = add_paragraph(doc, member, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # --- Section 1: Approach & Architecture ---
    add_heading(doc, "1. Project Overview & Methodology", level=1)
    add_paragraph(doc, "The objective of this project is to build a robust machine learning pipeline to predict hotel booking cancellations. We adopted a terminal-first workflow that separates model training and evaluation from the user-facing deployment (Streamlit app). This ensures a clean, production-ready architecture where evaluation metrics are completely honest.")
    
    add_heading(doc, "Data Processing & Engineering", level=2)
    add_paragraph(doc, "We engineered several high-signal features from the raw dataset:")
    doc.add_paragraph("• total_nights: Combination of weekend and weekday nights.", style='List Bullet')
    doc.add_paragraph("• family_booking: Indicates if the booking includes children or babies.", style='List Bullet')
    doc.add_paragraph("• previous_cancel_rate: The historical cancellation rate of the guest.", style='List Bullet')
    doc.add_paragraph("• room_match: Flag indicating if the assigned room matches the reserved room.", style='List Bullet')
    add_paragraph(doc, "Crucially, leakage columns like 'reservation_status' and 'reservation_status_date' were aggressively removed to ensure models generalize to real-world deployment scenarios where this data is unavailable at the time of booking.")

    add_heading(doc, "Pipeline Implementation Details", level=2)
    add_paragraph(doc, "Below is an excerpt of our `ModelTrainer` class in `hotel_app/ml/training.py`. It showcases the robust scikit-learn pipeline construction we used to ensure the preprocessor is bound tightly to the estimator, preventing data leakage during cross-validation.")
    code_snippet = '''def train_model(self, model_spec: BaseHotelModel, x_train: pd.DataFrame, y_train: pd.Series) -> Pipeline:
    # Build preprocessor (Scaling, One-Hot Encoding, Imputation) based on training data
    preprocessor = self.processor.build_preprocessor(x_train)
    
    # Construct a Sklearn Pipeline
    pipeline = model_spec.build_pipeline(preprocessor)
    
    # Fit the entire pipeline
    return pipeline.fit(x_train, y_train)'''
    add_code_block(doc, code_snippet)

    # --- Section 2: Models Trained ---
    doc.add_page_break()
    add_heading(doc, "2. Models Evaluated & Source Code", level=1)
    add_paragraph(doc, "We evaluated an exhaustive suite of 14 distinct machine learning algorithms. Each model is encapsulated in its own class, inheriting from `BaseHotelModel`. Below is the exact implementation code and functions used for each model in our pipeline.")
    doc.add_paragraph("\n")

    import inspect
    for name, model_class in MODEL_REGISTRY.items():
        doc.add_heading(name, level=2)
        try:
            source_code = inspect.getsource(model_class)
            add_code_block(doc, source_code.strip())
        except Exception as e:
            add_paragraph(doc, f"[Could not extract source code: {e}]", italic=True)
        doc.add_paragraph("\n")

    # --- Section 3: Metrics ---
    doc.add_page_break()
    add_heading(doc, "3. Evaluation Metrics (Holdout Split)", level=1)
    add_paragraph(doc, "Models were evaluated using an honest 30% holdout split. To provide complete transparency, the training time metrics below reflect both the benchmark phase and the full-dataset retraining phase required for deployment. The table below imports the live CSV metrics produced during the terminal run.")
    
    metrics_path = Path("artifacts/reports/holdout_summary.csv")
    if metrics_path.exists():
        df = pd.read_csv(metrics_path)
        # Select key columns
        df = df[['model', 'accuracy', 'f1', 'roc_auc', 'training_time_sec']].round(4)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Model'
        hdr_cells[1].text = 'Accuracy'
        hdr_cells[2].text = 'F1 Score'
        hdr_cells[3].text = 'ROC-AUC'
        hdr_cells[4].text = 'Training Time (s)'
        
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['model'])
            row_cells[1].text = str(row['accuracy'])
            row_cells[2].text = str(row['f1'])
            row_cells[3].text = str(row['roc_auc'])
            row_cells[4].text = str(row['training_time_sec'])
        doc.add_paragraph("\n")
    else:
        add_paragraph(doc, "[Metrics CSV not found. Please run training first.]", italic=True)

    add_heading(doc, "Performance & Timing Graphs", level=2)
    add_image(doc, "artifacts/plots/cross_validation_f1.png", width_inches=6.0, caption="Cross-Validation F1 Scores Across 5 Folds")
    add_image(doc, "artifacts/plots/timing_metrics.png", width_inches=6.0, caption="Model Training & Inference Time Analysis")

    # --- Section 4: Confusion Matrices ---
    doc.add_page_break()
    add_heading(doc, "4. Precision vs. Recall (Confusion Matrices)", level=1)
    add_paragraph(doc, "Below are the confusion matrices for several key models, illustrating the false-positive vs false-negative trade-offs made during prediction.")
    
    add_image(doc, "artifacts/plots/random_forest_confusion_matrix.png", width_inches=5.0, caption="Random Forest Confusion Matrix")
    add_image(doc, "artifacts/plots/xgboost_confusion_matrix.png", width_inches=5.0, caption="XGBoost Confusion Matrix")
    add_image(doc, "artifacts/plots/svm_confusion_matrix.png", width_inches=5.0, caption="SVM Confusion Matrix")

    # --- Section 5: Interpretability ---
    doc.add_page_break()
    add_heading(doc, "5. Explainability (SHAP Analysis)", level=1)
    add_paragraph(doc, "Machine Learning models should not act as black boxes. To provide transparency to the hotel management team, we implemented SHAP (SHapley Additive exPlanations) to explain the global decision-making process of the best performing models.")
    
    add_image(doc, "artifacts/plots/random_forest_shap_summary.png", width_inches=6.0, caption="Global SHAP Summary Plot: Features ranked by their impact on Cancellation Probability")
    
    add_paragraph(doc, "The SHAP plot visually demonstrates how features like 'deposit_type' (Non-Refundable deposits massively reduce cancellation risk) and 'lead_time' (longer lead times drastically increase cancellation risk) drive the predictions.")

    # Save Document
    doc.save(output_file)
    print(f"Detailed report successfully saved to {output_file}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Detailed Word Report")
    parser.add_argument("--members", nargs='+', help="List of team member names", default=None)
    args = parser.parse_args()
    
    build_report(team_members=args.members)
